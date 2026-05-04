"""
generate_preventivo.py — genera il preventivo .docx a partire dai dati di input.

USO:
    python3 generate_preventivo.py --input <modulo.xlsx> --output <preventivo.docx>
    python3 generate_preventivo.py --json '<dati_json>' --output <preventivo.docx>

STRUTTURA JSON INPUT (in alternativa all'xlsx):
{
  "numero": "2026-042",
  "data": "15/04/2026",               # opzionale, default = oggi
  "destinatario": "Studio Rossi Srl\\nVia Roma 1, 37051 Bovolone (VR)",
  "oggetto": "Rilievo capannone industriale",
  "validita_giorni": 30,               # opzionale, default 30
  "luogo": "Bovolone",                 # opzionale, default Bovolone
  "pagamento": "da concordare",        # opzionale, default "da concordare"
  "inserire_firma": true,              # opzionale, default true
  "note_aggiuntive": "",               # opzionale
  "voci": ["Rilievo laser scanner", "Restituzione nuvola punti CAD"],
  "totale": 1800.00                    # importo numerico senza IVA/CNG
}
"""
import argparse
import json
import os
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

# ----------------------------------------------------------
# DEFAULT
# ----------------------------------------------------------
DEFAULT_LUOGO = "Bovolone"
DEFAULT_PAGAMENTO = "da concordare"
DEFAULT_VALIDITA = 30
DEFAULT_FIRMA = True

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
TEMPLATE_PATH = SKILL_DIR / "assets" / "template_preventivo.docx"
FIRMA_PATH = SKILL_DIR / "assets" / "firma.png"

MESI_IT = {1: "gennaio", 2: "febbraio", 3: "marzo", 4: "aprile", 5: "maggio",
           6: "giugno", 7: "luglio", 8: "agosto", 9: "settembre",
           10: "ottobre", 11: "novembre", 12: "dicembre"}


# ----------------------------------------------------------
# INPUT: da Excel
# ----------------------------------------------------------
def parse_xlsx(xlsx_path):
    wb = load_workbook(xlsx_path, data_only=True)

    # --- FOGLIO DATI ---
    ws = wb["DATI"]
    dati = {}
    mapping = {
        "Numero preventivo": "numero",
        "Data": "data",
        "Destinatario": "destinatario",
        "Oggetto": "oggetto",
        "Validità (giorni)": "validita_giorni",
        "Luogo": "luogo",
        "Modalità pagamento": "pagamento",
        "Inserire firma (SI/NO)": "inserire_firma",
        "Note aggiuntive": "note_aggiuntive",
    }
    for row in ws.iter_rows(min_row=5, max_row=20, max_col=2, values_only=True):
        label, valore = row[0], row[1]
        if label and label in mapping:
            key = mapping[label]
            if valore is not None and str(valore).strip() != "":
                dati[key] = valore

    # --- FOGLIO VOCI ---
    ws2 = wb["VOCI"]
    voci_selezionate = []
    totale = None
    for row in ws2.iter_rows(min_row=5, values_only=True):
        spunta, attivita, col_c = row[0], row[1], row[2]
        if spunta is None and attivita is None:
            continue
        if attivita is None and col_c is not None and isinstance(col_c, (int, float)):
            totale = float(col_c)
            continue
        if attivita and isinstance(attivita, str):
            attivita_str = attivita.strip()
            if attivita_str.isupper() and (spunta is None or str(spunta).strip() == ""):
                continue
            if spunta is not None and str(spunta).strip() != "":
                voci_selezionate.append(attivita_str)

    dati["voci"] = voci_selezionate
    if totale is not None:
        dati["totale"] = totale

    return dati


# ----------------------------------------------------------
# NORMALIZZA DATI + applica default
# ----------------------------------------------------------
def normalize(dati):
    out = dict(dati)

    # Data
    if not out.get("data"):
        d = datetime.now()
        out["data"] = f"{d.day:02d}/{d.month:02d}/{d.year}"
    else:
        v = out["data"]
        if isinstance(v, datetime):
            out["data"] = f"{v.day:02d}/{v.month:02d}/{v.year}"
        else:
            out["data"] = str(v).strip()

    # Defaults
    out.setdefault("luogo", DEFAULT_LUOGO)
    out.setdefault("pagamento", DEFAULT_PAGAMENTO)
    out.setdefault("validita_giorni", DEFAULT_VALIDITA)
    out.setdefault("note_aggiuntive", "")

    # Firma (parsing SI/NO/bool)
    firma = out.get("inserire_firma", DEFAULT_FIRMA)
    if isinstance(firma, str):
        firma = firma.strip().upper() in ("SI", "SÌ", "YES", "Y", "TRUE", "1", "X")
    out["inserire_firma"] = bool(firma)

    # Validità giorni int
    try:
        out["validita_giorni"] = int(out["validita_giorni"])
    except (ValueError, TypeError):
        out["validita_giorni"] = DEFAULT_VALIDITA

    # Pulizia destinatario: rimuovi "Spett.le" se presente (viene aggiunto dal template)
    if out.get("destinatario"):
        dest = str(out["destinatario"]).strip()
        lowered = dest.lower()
        if lowered.startswith("spett.le") or lowered.startswith("spett.li") or lowered.startswith("spettabile"):
            for prefix in ("spett.le", "spett.li", "spettabile"):
                if lowered.startswith(prefix):
                    dest = dest[len(prefix):].lstrip(" \t:,.-—")
                    break
        out["destinatario"] = dest

    # Obbligatori
    if not out.get("destinatario"):
        raise ValueError("Campo obbligatorio mancante: destinatario")
    if not out.get("oggetto"):
        raise ValueError("Campo obbligatorio mancante: oggetto")
    if not out.get("numero"):
        raise ValueError("Campo obbligatorio mancante: numero preventivo")
    if not out.get("voci"):
        raise ValueError("Nessuna voce selezionata nel preventivo")
    if out.get("totale") is None:
        raise ValueError("Totale mancante")

    return out


# ----------------------------------------------------------
# HELPER docx
# ----------------------------------------------------------
def format_euro(value):
    """Formatta come '1.800,00' stile italiano."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    intero, dec = f"{v:,.2f}".split(".")
    intero = intero.replace(",", ".")
    return f"{intero},{dec}"


def set_cell_text(cell, text, bold=None, size=None, color=None, alignment=None, keep_style=True):
    """Sostituisce il testo di una cella preservando lo stile del primo run."""
    if keep_style and cell.paragraphs and cell.paragraphs[0].runs:
        ref_run = cell.paragraphs[0].runs[0]
        ref_para = cell.paragraphs[0]
        ref_font_name = ref_run.font.name
        ref_font_size = ref_run.font.size
        ref_bold = ref_run.bold
        ref_color = ref_run.font.color.rgb if ref_run.font.color and ref_run.font.color.rgb else None
        ref_align = ref_para.alignment
    else:
        ref_font_name = None
        ref_font_size = None
        ref_bold = None
        ref_color = None
        ref_align = None

    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = cell.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        elif ref_align is not None:
            p.alignment = ref_align
        r = p.add_run(line)
        if ref_font_name:
            r.font.name = ref_font_name
        if size is not None:
            r.font.size = Pt(size)
        elif ref_font_size:
            r.font.size = ref_font_size
        r.bold = bold if bold is not None else ref_bold
        if color is not None:
            r.font.color.rgb = color
        elif ref_color is not None:
            r.font.color.rgb = ref_color


def replace_in_paragraph(paragraph, replacements):
    """Sostituisce i placeholder in un paragrafo preservando lo stile."""
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, str(value))
    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


# ----------------------------------------------------------
# COMPOSIZIONE DOCUMENTO
# ----------------------------------------------------------
def build_docx(dati, output_path):
    doc = Document(str(TEMPLATE_PATH))

    data = dati["data"]
    numero = dati["numero"]
    destinatario = dati["destinatario"]
    oggetto = dati["oggetto"]
    pagamento = dati["pagamento"]
    validita = dati["validita_giorni"]
    luogo = dati["luogo"]
    voci = dati["voci"]
    totale = dati["totale"]
    note = dati.get("note_aggiuntive", "")
    inserire_firma = dati["inserire_firma"]

    # --- TABELLA 0: destinatario + header offerta ---
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 0), f"Spett.le\n{destinatario}")
    set_cell_text(t0.cell(0, 1), f"OFFERTA DI SERVIZIO\nOfferta N°: {numero}\nData: {data}")

    # --- PARAGRAFO OGGETTO ---
    for p in doc.paragraphs:
        if "[OGGETTO DEL PREVENTIVO]" in p.text:
            replace_in_paragraph(p, {"[OGGETTO DEL PREVENTIVO]": oggetto})
            break

    # --- TABELLA 1: voci e totale ---
    t1 = doc.tables[1]
    voice_rows_target = len(voci)
    existing_voice_rows = 3  # nel template

    while existing_voice_rows > voice_rows_target:
        row_to_remove = t1.rows[existing_voice_rows]
        row_to_remove._element.getparent().remove(row_to_remove._element)
        existing_voice_rows -= 1

    while existing_voice_rows < voice_rows_target:
        template_row = t1.rows[1]._element
        new_row = deepcopy(template_row)
        totale_row = t1.rows[-1]._element
        totale_row.addprevious(new_row)
        existing_voice_rows += 1

    for i, voce in enumerate(voci):
        row = t1.rows[1 + i]
        set_cell_text(row.cells[0], voce)
        set_cell_text(row.cells[1], "")

    totale_row = t1.rows[-1]
    set_cell_text(totale_row.cells[0], "TOTALE")
    set_cell_text(totale_row.cells[1], f"€ {format_euro(totale)} + IVA + CNG")

    # --- PARAGRAFI CONDIZIONI / VALIDITÀ / NOTA ---
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("Condizioni:"):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "Condizioni: "
                p.runs[0].bold = True
                new_run = p.add_run(pagamento)
                new_run.bold = False
                if p.runs[0].font.name:
                    new_run.font.name = p.runs[0].font.name
                if p.runs[0].font.size:
                    new_run.font.size = p.runs[0].font.size
            else:
                p.add_run("Condizioni: ").bold = True
                p.add_run(pagamento)
        elif txt.startswith("Validità:"):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "Validità: "
                p.runs[0].bold = True
                new_run = p.add_run(f"{validita} giorni dalla data dell'offerta.")
                new_run.bold = False
                if p.runs[0].font.name:
                    new_run.font.name = p.runs[0].font.name
                if p.runs[0].font.size:
                    new_run.font.size = p.runs[0].font.size
        elif txt.startswith("Nota:") and "laser scanner" in txt.lower():
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    # --- TABELLA 2: luogo/data firma + amministratore ---
    t2 = doc.tables[2]
    set_cell_text(t2.cell(0, 0),
                  f"{luogo}, lì {data}\n\nFirma per accettazione {numero}\n\n\n___________________________")

    cell_firma = t2.cell(0, 1)
    for p in list(cell_firma.paragraphs):
        p._element.getparent().remove(p._element)

    firma_esiste = FIRMA_PATH.exists() and FIRMA_PATH.stat().st_size > 0
    if inserire_firma and firma_esiste:
        p_firma = cell_firma.add_paragraph()
        p_firma.alignment = 2  # right
        r = p_firma.add_run()
        try:
            r.add_picture(str(FIRMA_PATH), height=Cm(1.8))
        except Exception as e:
            print(f"WARN: impossibile inserire firma: {e}", file=sys.stderr)
    elif inserire_firma and not firma_esiste:
        print(f"WARN: file firma non trovato in {FIRMA_PATH}. Preventivo generato senza firma.", file=sys.stderr)

    p_amm = cell_firma.add_paragraph()
    p_amm.alignment = 2
    r_amm = p_amm.add_run("Amministratore")
    r_amm.font.size = Pt(10)

    p_nome = cell_firma.add_paragraph()
    p_nome.alignment = 2
    r_nome = p_nome.add_run("LEONARDO MASSAFRA")
    r_nome.bold = True
    r_nome.font.size = Pt(11)

    doc.save(str(output_path))
    return output_path


# ----------------------------------------------------------
# MAIN
# ----------------------------------------------------------
def safe_filename(s):
    return "".join(c if c.isalnum() or c in "-_ " else "_" for c in s).strip().replace(" ", "_")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", help="Percorso al modulo .xlsx")
    parser.add_argument("--json", help="Dati in formato JSON inline")
    parser.add_argument("--output", help="Percorso del .docx di output")
    parser.add_argument("--output-dir", help="Cartella di output (nome generato automaticamente)")
    args = parser.parse_args()

    if args.json:
        dati = json.loads(args.json)
    elif args.input:
        dati = parse_xlsx(args.input)
    else:
        print("ERRORE: specificare --input o --json", file=sys.stderr)
        sys.exit(1)

    dati = normalize(dati)

    if args.output:
        output_path = Path(args.output)
    else:
        output_dir = Path(args.output_dir) if args.output_dir else Path.cwd()
        output_dir.mkdir(parents=True, exist_ok=True)
        dest_short = dati["destinatario"].split("\n")[0][:40]
        filename = f"Preventivo_{safe_filename(dati['numero'])}_{safe_filename(dest_short)}.docx"
        output_path = output_dir / filename

    build_docx(dati, output_path)
    print(f"✓ Preventivo generato: {output_path}")


if __name__ == "__main__":
    main()
